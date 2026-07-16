from pathlib import Path
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "AutoDoctor_TZ_MVP.md"
TARGET = ROOT / "AutoDoctor_TZ_MVP.docx"
TEMPORARY = ROOT / "AutoDoctor_TZ_MVP.docx.tmp"


def set_font(run, name="Arial", size=10.5, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold


def add_inline(paragraph, text):
    text = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        lambda match: f"{match.group(1)} ({match.group(2)})",
        text,
    )
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", 9.5)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def configure(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for level, size in {1: 18, 2: 14, 3: 11.5}.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.keep_with_next = True

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def export():
    document = Document()
    configure(document)
    document.core_properties.title = "AutoDoctor — техническое задание на MVP"
    document.core_properties.subject = "Закрытый MVP-пилот в Беларуси"

    in_code = False
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line and not in_code:
            continue
        if in_code:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line or " ")
            set_font(run, "Consolas", 8.5)
            paragraph.paragraph_format.left_indent = Cm(0.8)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph = document.add_heading(level=level - 1)
            add_inline(paragraph, heading.group(2))
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
        elif numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, line)

    document.save(TEMPORARY)
    os.replace(TEMPORARY, TARGET)


if __name__ == "__main__":
    export()
